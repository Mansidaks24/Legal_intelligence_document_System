"""
parser.py  —  IMPROVED VERSION
--------------------------------
User Stories covered:
  US-03 : Text Extraction (PDF + DOCX)

Additional features added (add to Excel as extras):
  EXTRA-01 : DOCX file support (python-docx)
  EXTRA-02 : Table text extraction from DOCX
  EXTRA-03 : Page-level metadata tracking (page count)
  EXTRA-04 : Character cleaning (removes junk/non-printable chars)
"""

import logging
import re
from pathlib import Path

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)

# ── Max file size allowed: 10 MB ──────────────────────────────────────
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB


def extract_text(file_path: str) -> dict:
    """
    Extract text from a PDF or DOCX file.

    Returns a dict with:
      - text          : full extracted text string
      - page_count    : number of pages (PDF) or sections (DOCX)
      - file_type     : 'pdf' or 'docx'
      - char_count    : total characters extracted

    Raises:
      FileNotFoundError : file path does not exist
      ValueError        : unsupported format or empty content
      RuntimeError      : parsing failure
    """
    path = Path(file_path)

    # ── File existence check ─────────────────────────────────────────
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    # ── File size check (EXTRA-01 security) ─────────────────────────
    file_size = path.stat().st_size
    if file_size == 0:
        raise ValueError("Uploaded file is empty.")
    if file_size > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"File too large ({file_size // (1024*1024)} MB). "
            f"Maximum allowed size is 10 MB."
        )

    # ── Route to correct extractor by file type ──────────────────────
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    elif suffix == ".docx":
        return _extract_docx(path)
    else:
        raise ValueError(
            f"Unsupported file type: '{suffix}'. "
            "Only .pdf and .docx files are accepted."
        )


# ─────────────────────────────────────────────────────────────────────
# PDF Extractor
# ─────────────────────────────────────────────────────────────────────
def _extract_pdf(path: Path) -> dict:
    """Extract text page by page from a digitally-created PDF."""
    extracted_pages = []

    try:
        with pdfplumber.open(str(path)) as pdf:
            total_pages = len(pdf.pages)
            logger.info(f"Opening PDF: '{path.name}' ({total_pages} pages)")

            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    cleaned = _clean_text(page_text)
                    extracted_pages.append(cleaned)
                    logger.debug(f"  Page {page_num}/{total_pages}: {len(cleaned)} chars")
                else:
                    logger.warning(
                        f"  Page {page_num}/{total_pages}: no text "
                        "(possibly a scanned image page)"
                    )

    except Exception as exc:
        logger.error(f"PDF parse error '{path.name}': {exc}")
        raise RuntimeError(f"PDF parsing failed: {exc}") from exc

    if not extracted_pages:
        raise ValueError(
            "No extractable text found in the PDF. "
            "The document may be scanned — OCR support can be added later."
        )

    full_text = "\n\n".join(extracted_pages)
    logger.info(f"PDF extraction complete: {len(full_text)} characters, {len(extracted_pages)} pages")

    return {
        "text": full_text,
        "page_count": len(extracted_pages),
        "file_type": "pdf",
        "char_count": len(full_text),
    }


# ─────────────────────────────────────────────────────────────────────
# DOCX Extractor  (EXTRA-01, EXTRA-02)
# ─────────────────────────────────────────────────────────────────────
def _extract_docx(path: Path) -> dict:
    """
    Extract text from a Word (.docx) document.
    Captures both paragraph text AND table cell text (EXTRA-02).
    """
    extracted_parts = []

    try:
        doc = Document(str(path))
        logger.info(f"Opening DOCX: '{path.name}'")

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                extracted_parts.append(text)

        # Extract table contents (EXTRA-02)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    extracted_parts.append(row_text)

    except Exception as exc:
        logger.error(f"DOCX parse error '{path.name}': {exc}")
        raise RuntimeError(f"DOCX parsing failed: {exc}") from exc

    if not extracted_parts:
        raise ValueError("No extractable text found in the DOCX file.")

    full_text = _clean_text("\n\n".join(extracted_parts))
    logger.info(f"DOCX extraction complete: {len(full_text)} characters")

    return {
        "text": full_text,
        "page_count": len(extracted_parts),   # paragraph count for DOCX
        "file_type": "docx",
        "char_count": len(full_text),
    }


# ─────────────────────────────────────────────────────────────────────
# Text Cleaner  (EXTRA-04)
# ─────────────────────────────────────────────────────────────────────
def _clean_text(text: str) -> str:
    """
    Remove junk characters from extracted text.
    Keeps: letters, digits, punctuation, whitespace.
    Removes: non-printable control characters, null bytes.
    """
    # Remove null bytes and control characters (except \n and \t)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    # Collapse 3+ consecutive newlines into 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse multiple spaces into one
    text = re.sub(r" {2,}", " ", text)
    return text.strip()